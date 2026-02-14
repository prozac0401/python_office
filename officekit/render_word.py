from __future__ import annotations

from pathlib import Path

import pandas as pd

from .column_alias import add_alias_keys, resolve_existing_column
from .io_excel import read_excel_to_df
from .utils import ensure_dir, safe_filename


def _next_available_path(path: Path) -> Path:
    """파일명이 충돌하면 `_N`을 붙여 사용 가능한 경로를 반환한다."""
    if not path.exists():
        return path
    stem = path.stem
    suffix = path.suffix
    idx = 2
    while True:
        candidate = path.with_name(f"{stem}_{idx}{suffix}")
        if not candidate.exists():
            return candidate
        idx += 1


def _normalize_ctx(df: pd.DataFrame, row: pd.Series, include_aliases: bool = False) -> dict[str, object]:
    """DataFrame의 한 행을 템플릿 컨텍스트 딕셔너리로 변환한다."""
    ctx: dict[str, object] = {}
    for col in df.columns:
        val = row[col]
        if pd.isna(val):
            ctx[col] = ""
        else:
            ctx[col] = val

    # 자주 쓰는 파생 필드
    qty_key = resolve_existing_column(ctx.keys(), "Qty")
    unit_key = resolve_existing_column(ctx.keys(), "UnitPrice")
    total_key = resolve_existing_column(ctx.keys(), "Total")
    if total_key is None and qty_key is not None and unit_key is not None:
        total_name = "합계" if (qty_key == "수량" or unit_key == "단가") else "Total"
        try:
            ctx[total_name] = float(ctx[qty_key]) * float(ctx[unit_key])
        except Exception:
            pass

    return add_alias_keys(ctx) if include_aliases else ctx


def _replace_in_paragraph_runs(paragraph, mapping: dict[str, str]) -> None:
    # 최대한 서식을 유지하기 위해 run 단위로 치환한다.
    # 주의: 플레이스홀더가 여러 run으로 분리되면 치환되지 않을 수 있다.
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        for k, v in mapping.items():
            text = text.replace(f"{{{{ {k} }}}}", v).replace(f"{{{{{k}}}}}", v)
        run.text = text


def _replace_in_doc(doc, mapping: dict[str, str]) -> None:
    # 본문 문단
    for p in doc.paragraphs:
        _replace_in_paragraph_runs(p, mapping)

    # 표
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_runs(p, mapping)

    # 머리글/바닥글(일반적인 경우)
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_in_paragraph_runs(p, mapping)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph_runs(p, mapping)

        for p in section.footer.paragraphs:
            _replace_in_paragraph_runs(p, mapping)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _replace_in_paragraph_runs(p, mapping)


def _render_with_template(template_docx: Path, ctx: dict[str, object], out_path: Path) -> None:
    """DOCX 템플릿을 렌더링한다.

    우선순위:
    1) `docxtpl`이 설치되어 있으면 우선 사용(더 견고한 템플릿 처리)
    2) 없으면 python-docx로 단순 `{{var}}` 치환 수행
    """
    # 1) docxtpl을 사용할 수 있으면 우선 사용한다.
    # 설치되어 있는데 렌더링이 실패하면 예외를 그대로 노출한다.
    try:
        from docxtpl import DocxTemplate  # type: ignore
    except ImportError:
        DocxTemplate = None  # type: ignore[assignment]

    if DocxTemplate is not None:
        tpl = DocxTemplate(str(template_docx))
        tpl.render(ctx)
        tpl.save(str(out_path))
        return

    # 2) 단순 치환(추가 의존성 없음)
    from docx import Document  # 지연 import

    doc = Document(str(template_docx))
    mapping = {k: "" if v is None else str(v) for k, v in ctx.items()}
    _replace_in_doc(doc, mapping)
    doc.save(str(out_path))


def _render_simple_docx(ctx: dict[str, object], out_path: Path, title_template: str) -> None:
    """템플릿이 없을 때 사용하는 python-docx 기반 기본 렌더러."""
    from docx import Document  # 지연 import

    doc = Document()
    fmt_ctx = add_alias_keys(ctx)
    try:
        title = title_template.format(**{k: str(v) for k, v in fmt_ctx.items()})
    except Exception:
        title = "행 문서"
    doc.add_heading(title, level=1)

    doc.add_paragraph("이 문서는 Excel 행 데이터를 기반으로 자동 생성되었습니다.")

    keys = list(ctx.keys())
    preferred_order = [
        k for k in ("아이디", "이름", "부서", "분류", "주문일자", "이메일", "수량", "단가", "합계", "ID", "Name", "Department", "Category", "OrderDate", "Email", "Qty", "UnitPrice", "Total")
        if k in ctx
    ]
    rest = [k for k in keys if k not in preferred_order]
    ordered = preferred_order + rest

    table = doc.add_table(rows=len(ordered), cols=2)
    table.style = "Table Grid"
    for i, k in enumerate(ordered):
        table.cell(i, 0).text = str(k)
        table.cell(i, 1).text = str(ctx.get(k, ""))

    doc.save(str(out_path))


def create_word_docs_from_excel(
    input_xlsx: str | Path,
    output_dir: str | Path,
    sheet_name: int | str = 0,
    template_docx: str | Path | None = None,
    filename_cols: tuple[str, ...] = ("ID", "Name"),
    title_template: str = "주문 요약 - {ID} / {Name}",
) -> list[Path]:
    """Excel 시트의 각 행마다 DOCX 파일 1개를 생성한다.

    - `template_docx`가 있으면 템플릿 렌더링을 수행한다.
    - 없으면 python-docx로 기본 문서를 생성한다.

    템플릿 변수:
    - Excel 컬럼명은 그대로 변수로 사용 가능(예: `{{ID}}`, `{{Name}}`, `{{아이디}}`, `{{이름}}`)
    파생 변수:
    - `Qty/UnitPrice` 또는 `수량/단가`가 있으면 `Total/합계`를 자동 계산한다.
    """
    df = read_excel_to_df(input_xlsx, sheet_name=sheet_name)
    out_dir = ensure_dir(output_dir)
    out_paths: list[Path] = []

    tpl_path: Path | None = Path(template_docx) if template_docx else None
    if tpl_path is not None and not tpl_path.exists():
        raise FileNotFoundError(f"템플릿 파일을 찾을 수 없습니다: {tpl_path}")

    for _, row in df.iterrows():
        raw_ctx = _normalize_ctx(df, row, include_aliases=False)
        ctx = add_alias_keys(raw_ctx)

        parts: list[str] = []
        for col in filename_cols:
            if col in ctx and str(ctx[col]).strip():
                parts.append(str(ctx[col]))
        fname = safe_filename("_".join(parts) if parts else "row") + ".docx"
        out_path = _next_available_path(out_dir / fname)

        if tpl_path is not None:
            _render_with_template(tpl_path, ctx, out_path)
        else:
            _render_simple_docx(raw_ctx, out_path, title_template=title_template)

        out_paths.append(out_path)

    return out_paths
